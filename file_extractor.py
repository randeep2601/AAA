import os
import zipfile
import tempfile
from typing import List, Optional
import logging

# Import libraries for different file types
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import olefile
    import struct
    DOC_AVAILABLE = True
except ImportError:
    DOC_AVAILABLE = False

class FileExtractor:
    def __init__(self):
        self.max_file_size = 100 * 1024 * 1024  # 100MB
        self.max_pages = 1000  # Maximum pages for PDF processing
        
    def extract_text(self, file_path: str) -> str:
        """Extract text from various file formats"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size:
                raise Exception(f"File too large: {file_size / (1024*1024):.1f}MB. Maximum allowed: {self.max_file_size / (1024*1024)}MB")
            
            if file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext in ['.html', '.htm']:
                return self._extract_from_html(file_path)
            elif file_ext == '.zip':
                return self._extract_from_zip(file_path)
            elif file_ext == '.docx':
                return self._extract_from_docx(file_path)
            elif file_ext == '.doc':
                return self._extract_from_doc(file_path)
            elif file_ext == '.txt':
                return self._extract_from_txt(file_path)
            else:
                raise Exception(f"Unsupported file type: {file_ext}")
                
        except Exception as e:
            raise Exception(f"Failed to extract text: {str(e)}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        if not PDF_AVAILABLE:
            raise Exception("PDF processing libraries not available. Please install PyPDF2 and pdfplumber.")
        
        text_content = []
        
        try:
            # Try with pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                total_pages = min(len(pdf.pages), self.max_pages)
                
                for page_num in range(total_pages):
                    page = pdf.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text:
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}\n")
                
                if len(pdf.pages) > self.max_pages:
                    text_content.append(f"\n--- NOTE: Only first {self.max_pages} pages processed ---\n")
                    
        except Exception as e:
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    total_pages = min(len(pdf_reader.pages), self.max_pages)
                    
                    for page_num in range(total_pages):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        
                        if page_text:
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}\n")
                    
                    if len(pdf_reader.pages) > self.max_pages:
                        text_content.append(f"\n--- NOTE: Only first {self.max_pages} pages processed ---\n")
                        
            except Exception as e2:
                raise Exception(f"Could not extract text from PDF: {str(e2)}")
        
        if not text_content:
            raise Exception("No text could be extracted from the PDF file")
        
        return '\n'.join(text_content)
    
    def _extract_from_html(self, file_path: str) -> str:
        """Extract text from HTML files"""
        if not HTML_AVAILABLE:
            raise Exception("HTML processing library not available. Please install beautifulsoup4.")
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                html_content = file.read()
            
            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except Exception as e:
            raise Exception(f"Could not extract text from HTML file: {str(e)}")
    
    def _extract_from_zip(self, file_path: str) -> str:
        """Extract text from ZIP files containing supported documents"""
        try:
            text_content = []
            
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # Get list of files in the ZIP
                file_list = zip_file.namelist()
                
                # Filter for supported file types
                supported_files = []
                for filename in file_list:
                    file_ext = os.path.splitext(filename)[1].lower()
                    if file_ext in ['.pdf', '.html', '.htm', '.docx', '.doc', '.txt']:
                        supported_files.append(filename)
                
                if not supported_files:
                    raise Exception("No supported file types found in ZIP archive")
                
                # Extract and process each supported file
                with tempfile.TemporaryDirectory() as temp_dir:
                    for filename in supported_files:
                        try:
                            # Extract file to temporary directory
                            zip_file.extract(filename, temp_dir)
                            extracted_path = os.path.join(temp_dir, filename)
                            
                            # Extract text from the file
                            file_text = self.extract_text(extracted_path)
                            
                            if file_text.strip():
                                text_content.append(f"=== File: {filename} ===\n{file_text}\n")
                                
                        except Exception as e:
                            text_content.append(f"=== File: {filename} ===\nError extracting: {str(e)}\n")
            
            if not text_content:
                raise Exception("No text could be extracted from any files in the ZIP archive")
            
            return '\n'.join(text_content)
            
        except Exception as e:
            raise Exception(f"Could not extract text from ZIP file: {str(e)}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        if not DOCX_AVAILABLE:
            raise Exception("DOCX processing library not available. Please install python-docx.")
        
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            if not text_content:
                raise Exception("No text found in the DOCX file")
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            raise Exception(f"Could not extract text from DOCX file: {str(e)}")
    
    def _extract_from_doc(self, file_path: str) -> str:
        """Extract text from DOC files (basic extraction)"""
        try:
            # This is a basic implementation for older DOC files
            # For better results, consider using additional libraries like python-docx2txt
            
            with open(file_path, 'rb') as file:
                content = file.read()
            
            # Try to extract readable text (very basic approach)
            text_content = []
            current_text = ""
            
            for byte in content:
                if 32 <= byte <= 126:  # Printable ASCII characters
                    current_text += chr(byte)
                else:
                    if len(current_text) > 3:  # Only add substantial text chunks
                        text_content.append(current_text)
                    current_text = ""
            
            if current_text and len(current_text) > 3:
                text_content.append(current_text)
            
            # Join and clean the text
            extracted_text = ' '.join(text_content)
            
            # Basic cleanup
            import re
            extracted_text = re.sub(r'\s+', ' ', extracted_text)
            extracted_text = re.sub(r'[^\w\s.,!?;:()-]', '', extracted_text)
            
            if not extracted_text.strip():
                raise Exception("No readable text found in the DOC file")
            
            return extracted_text.strip()
            
        except Exception as e:
            raise Exception(f"Could not extract text from DOC file: {str(e)}. Consider converting to DOCX format for better results.")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    
                    if content.strip():
                        return content
                        
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            if not content.strip():
                raise Exception("Text file appears to be empty")
            
            return content
            
        except Exception as e:
            raise Exception(f"Could not extract text from TXT file: {str(e)}")
    
    def get_file_info(self, file_path: str) -> dict:
        """Get information about the file"""
        try:
            file_size = os.path.getsize(file_path)
            file_ext = os.path.splitext(file_path)[1].lower()
            
            info = {
                'size_bytes': file_size,
                'size_mb': round(file_size / (1024 * 1024), 2),
                'extension': file_ext,
                'supported': file_ext in ['.pdf', '.html', '.htm', '.zip', '.doc', '.docx', '.txt']
            }
            
            # Add format-specific info
            if file_ext == '.pdf' and PDF_AVAILABLE:
                try:
                    with pdfplumber.open(file_path) as pdf:
                        info['pages'] = len(pdf.pages)
                        info['will_process_pages'] = min(len(pdf.pages), self.max_pages)
                except:
                    info['pages'] = 'Unknown'
            
            return info
            
        except Exception as e:
            return {'error': str(e)}